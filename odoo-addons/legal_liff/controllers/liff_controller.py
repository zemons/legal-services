import io
import json
import logging
import os
import re
import subprocess
import tempfile
import threading

from odoo import http, fields
from odoo.http import request

_logger = logging.getLogger(__name__)

CASE_TYPES = [
    ('civil', 'แพ่ง'),
    ('criminal', 'อาญา'),
    ('family', 'ครอบครัว'),
    ('inheritance', 'มรดก'),
    ('land', 'ที่ดิน'),
    ('labor', 'แรงงาน'),
    ('admin', 'ปกครอง'),
    ('ip', 'ทรัพย์สินทางปัญญา'),
    ('construction', 'ก่อสร้าง'),
    ('insurance', 'ประกันภัย'),
    ('document', 'งานเอกสาร'),
    ('consult', 'ที่ปรึกษา'),
    ('debt', 'บังคับคดี/ติดตามหนี้'),
]

CASE_STATUS_LABELS = dict([
    ('intake', 'รับเรื่อง'),
    ('review', 'ทนายกำลังตรวจสอบ'),
    ('in_progress', 'กำลังดำเนินการ'),
    ('court_pending', 'รอนัดศาล'),
    ('court_ongoing', 'อยู่ระหว่างพิจารณา'),
    ('settled', 'ยุติ/ไกล่เกลี่ย'),
    ('closed_won', 'ปิดคดี - ชนะ'),
    ('closed_lost', 'ปิดคดี - แพ้'),
    ('closed_other', 'ปิดคดี - อื่นๆ'),
])

CASE_STATUS_COLORS = {
    'intake': '#3498db',
    'review': '#f39c12',
    'in_progress': '#2ecc71',
    'court_pending': '#e67e22',
    'court_ongoing': '#9b59b6',
    'settled': '#1abc9c',
    'closed_won': '#27ae60',
    'closed_lost': '#e74c3c',
    'closed_other': '#95a5a6',
}

PURPOSE_LABELS = dict([
    ('hearing', 'นัดพิจารณา'),
    ('witness', 'นัดสืบพยาน'),
    ('mediation', 'นัดไกล่เกลี่ย'),
    ('judgment', 'นัดฟังคำพิพากษา'),
    ('other', 'อื่นๆ'),
])


class LiffController(http.Controller):

    # ── helpers ──────────────────────────────────────────────

    def _get_partner_from_line(self, line_user_id):
        """Returns (partner recordset, line_role string) or (empty, None)."""
        if not line_user_id:
            return request.env['res.partner'].sudo(), None
        partner = request.env['res.partner'].sudo().search(
            [('line_user_id', '=', line_user_id)], limit=1)
        return partner, partner.line_role if partner else None

    def _get_template_steps(self, template):
        """Build step-based field structure for guided interview.

        Returns: [
            {"step": 1, "label": "ข้อมูลคู่สัญญา", "fields": [...]},
            {"step": 2, "label": "เงื่อนไขสัญญา", "fields": [...]},
        ]
        """
        field_records = template.field_ids.sorted(lambda f: (f.step, f.sequence))
        if not field_records:
            return []

        steps_dict = {}
        for f in field_records:
            step_num = f.step or 1
            if step_num not in steps_dict:
                steps_dict[step_num] = {
                    'step': step_num,
                    'label': f.step_label or f'ขั้นตอนที่ {step_num}',
                    'fields': [],
                }
            field_data = {
                'name': f.name,
                'label': f.label,
                'type': f.field_type,
                'required': f.required,
                'options': f.options or '',
                'default_value': f.default_value or '',
                'show_when': f.show_when or '',
                'help_text': f.help_text or '',
                'placeholder': f.placeholder or '',
            }
            if f.field_type == 'repeating' and f.repeating_fields_json:
                try:
                    field_data['repeating_fields'] = json.loads(f.repeating_fields_json)
                except (json.JSONDecodeError, TypeError):
                    field_data['repeating_fields'] = []
                field_data['repeating_min'] = f.repeating_min
                field_data['repeating_max'] = f.repeating_max
            steps_dict[step_num]['fields'].append(field_data)

        # Use the step_label from the first field of each step
        for step_num, step_data in steps_dict.items():
            first_with_label = next(
                (f for f in field_records
                 if (f.step or 1) == step_num and f.step_label),
                None)
            if first_with_label:
                step_data['label'] = first_with_label.step_label

        return sorted(steps_dict.values(), key=lambda s: s['step'])

    def _get_liff_id(self):
        return request.env['ir.config_parameter'].sudo().get_param(
            'line_integration.liff_id', '0000000000-xxxxxxxx')

    def _send_line_push(self, line_user_id, lead, name):
        try:
            import requests as req
            access_token = request.env['ir.config_parameter'].sudo().get_param(
                'line_integration.channel_access_token', '')
            if not access_token:
                return
            case_label = dict(CASE_TYPES).get(lead.case_type, '')
            resp = req.post(
                'https://api.line.me/v2/bot/message/push',
                headers={
                    'Authorization': f'Bearer {access_token}',
                    'Content-Type': 'application/json',
                },
                json={
                    'to': line_user_id,
                    'messages': [{
                        'type': 'text',
                        'text': (
                            f"รับเรื่องของคุณแล้วค่ะ (เลขที่ {lead.id})\n"
                            f"ประเภท: {case_label}\n\n"
                            "ทนายความจะตรวจสอบและติดต่อกลับโดยเร็วค่ะ"
                        ),
                    }],
                },
            )
            if resp.status_code != 200:
                _logger.error('LINE push failed: %s', resp.text)
        except Exception as e:
            _logger.error('LINE push error: %s', e)

    # ── /liff/intake ─────────────────────────────────────────

    @http.route('/liff/intake', type='http', auth='public', methods=['GET'], csrf=False)
    def liff_intake(self, **kwargs):
        return request.render('legal_liff.liff_intake_page', {
            'case_types': CASE_TYPES,
            'error': '',
            'values': {},
            'liff_id': self._get_liff_id(),
        })

    @http.route('/liff/intake/submit', type='http', auth='public', methods=['POST'], csrf=True)
    def liff_intake_submit(self, **post):
        name = post.get('name', '').strip()
        phone = post.get('phone', '').strip()
        case_type = post.get('case_type', '')
        description = post.get('description', '').strip()
        line_user_id = post.get('line_user_id', '').strip()

        if not name or not case_type or not description:
            return request.render('legal_liff.liff_intake_page', {
                'case_types': CASE_TYPES,
                'error': 'กรุณากรอกข้อมูลให้ครบ: ชื่อ, ประเภทคดี, รายละเอียด',
                'values': post,
                'liff_id': self._get_liff_id(),
            })

        # Find or create partner
        Partner = request.env['res.partner'].sudo()
        partner = None
        if line_user_id:
            partner = Partner.search([('line_user_id', '=', line_user_id)], limit=1)
        if not partner and phone:
            partner = Partner.search([('phone', '=', phone)], limit=1)
        if not partner:
            partner = Partner.create({
                'name': name,
                'phone': phone,
                'line_user_id': line_user_id or False,
            })
        elif not partner.phone and phone:
            partner.write({'phone': phone})

        # Create CRM Lead
        case_type_labels = dict(CASE_TYPES)
        lead_name = f"{case_type_labels.get(case_type, case_type)} - {name}"
        lead = request.env['crm.lead'].sudo().create({
            'name': lead_name,
            'partner_id': partner.id,
            'phone': phone,
            'case_type': case_type,
            'case_status': 'intake',
            'description': description,
            'type': 'opportunity',
        })

        # Create initial timeline event
        request.env['legal.case.event'].sudo().create({
            'lead_id': lead.id,
            'event_type': 'status_change',
            'title': 'รับเรื่องใหม่',
            'description': f'ลูกค้า {name} ส่งเรื่องผ่าน LIFF',
        })

        _logger.info('LIFF intake: created lead #%s for %s', lead.id, name)

        if line_user_id:
            self._send_line_push(line_user_id, lead, name)

        return request.render('legal_liff.liff_intake_success', {
            'lead': lead,
            'name': name,
            'liff_id': self._get_liff_id(),
        })

    # ── /liff/status/<id> ────────────────────────────────────

    @http.route('/liff/status/<int:lead_id>', type='http', auth='public', methods=['GET'], csrf=False)
    def liff_status(self, lead_id, **kwargs):
        lead = request.env['crm.lead'].sudo().browse(lead_id)
        if not lead.exists():
            return request.render('legal_liff.liff_error_page', {
                'message': 'ไม่พบข้อมูลคดี',
                'liff_id': self._get_liff_id(),
            })

        events = request.env['legal.case.event'].sudo().search(
            [('lead_id', '=', lead_id)], order='date desc, id desc')

        now = fields.Datetime.now()
        upcoming_courts = lead.court_date_ids.filtered(
            lambda d: d.date_time and d.date_time >= now
        ).sorted('date_time')

        past_courts = lead.court_date_ids.filtered(
            lambda d: d.date_time and d.date_time < now
        ).sorted('date_time', reverse=True)

        case_type_labels = dict(CASE_TYPES)

        return request.render('legal_liff.liff_status_page', {
            'lead': lead,
            'events': events,
            'upcoming_courts': upcoming_courts,
            'past_courts': past_courts,
            'status_label': CASE_STATUS_LABELS.get(lead.case_status, lead.case_status or '-'),
            'status_color': CASE_STATUS_COLORS.get(lead.case_status, '#95a5a6'),
            'case_type_label': case_type_labels.get(lead.case_type, lead.case_type or '-'),
            'purpose_labels': PURPOSE_LABELS,
            'liff_id': self._get_liff_id(),
        })

    # ── /liff/cases (Lawyer Dashboard) ───────────────────────

    @http.route('/liff/cases', type='http', auth='public', methods=['GET'], csrf=False)
    def liff_cases(self, **kwargs):
        return request.render('legal_liff.liff_cases_page', {
            'liff_id': self._get_liff_id(),
            'status_labels': json.dumps(CASE_STATUS_LABELS, ensure_ascii=False),
            'status_colors': json.dumps(CASE_STATUS_COLORS),
            'case_type_labels': json.dumps(dict(CASE_TYPES), ensure_ascii=False),
        })

    @http.route('/liff/cases/data', type='json', auth='public', methods=['POST'], csrf=False)
    def liff_cases_data(self, **kwargs):
        line_user_id = kwargs.get('line_user_id', '')
        status_filter = kwargs.get('status_filter', 'all')

        partner, role = self._get_partner_from_line(line_user_id)
        if not partner or role != 'lawyer':
            return {'error': 'unauthorized', 'message': 'เฉพาะทนายความเท่านั้น'}

        user = request.env['res.users'].sudo().search(
            [('partner_id', '=', partner.id)], limit=1)
        if not user:
            return {'error': 'no_user', 'message': 'ไม่พบบัญชีผู้ใช้'}

        domain = [
            '|',
            ('user_id', '=', user.id),
            ('collaborator_ids', 'in', user.id),
        ]

        if status_filter == 'active':
            domain.append(('case_status', 'not in',
                           ['closed_won', 'closed_lost', 'closed_other', 'settled']))
        elif status_filter == 'court':
            domain.append(('case_status', 'in', ['court_pending', 'court_ongoing']))
        elif status_filter == 'closed':
            domain.append(('case_status', 'in',
                           ['closed_won', 'closed_lost', 'closed_other', 'settled']))

        leads = request.env['crm.lead'].sudo().search(domain, order='write_date desc')

        now = fields.Datetime.now()
        cases = []
        for lead in leads:
            next_court = lead.court_date_ids.filtered(
                lambda d: d.date_time and d.date_time >= now
            ).sorted('date_time')[:1]

            cases.append({
                'id': lead.id,
                'name': lead.name,
                'case_type': lead.case_type,
                'case_status': lead.case_status,
                'partner_name': lead.partner_id.name or '-',
                'next_court_date': str(next_court.date_time) if next_court else '',
                'next_court_purpose': next_court.purpose if next_court else '',
            })

        return {
            'cases': cases,
            'lawyer_name': partner.name,
            'total': len(cases),
        }

    # ── /liff/document/create ─────────────────────────────────

    @http.route('/liff/document/create', type='http', auth='public', methods=['GET'], csrf=False)
    def liff_document_create(self, **kwargs):
        return request.render('legal_liff.liff_document_create_page', {
            'liff_id': self._get_liff_id(),
        })

    @http.route('/liff/document/test', type='http', auth='public', methods=['GET'], csrf=False)
    def liff_document_test(self, **kwargs):
        """Standalone test page — no LIFF SDK dependency."""
        html = """<!DOCTYPE html>
<html><head>
<meta charset="utf-8"/>
<meta name="viewport" content="width=device-width, initial-scale=1"/>
<title>Test Document Create</title>
<link href="https://fonts.googleapis.com/css2?family=Noto+Sans+Thai:wght@400;600;700&display=swap" rel="stylesheet"/>
<style>
body { font-family: 'Noto Sans Thai', sans-serif; background: #f5f6fa; margin: 0; }
.c { max-width: 480px; margin: 0 auto; padding: 16px; background: #fff; min-height: 100vh; }
h2 { text-align: center; font-size: 22px; margin: 20px 0 10px; }
.sub { text-align: center; color: #7f8c8d; font-size: 14px; margin-bottom: 16px; }
label { display: block; font-size: 14px; font-weight: 600; margin-bottom: 6px; }
select, input, textarea { width: 100%; padding: 12px; border: 1.5px solid #dfe6e9; border-radius: 10px; font-size: 16px; font-family: inherit; background: #f8f9fa; box-sizing: border-box; margin-bottom: 16px; -webkit-appearance: none; }
select:focus, input:focus, textarea:focus { border-color: #2ECC71; outline: none; background: #fff; }
.btn { display: block; width: 100%; padding: 14px; background: #2ECC71; color: #fff; border: none; border-radius: 10px; font-size: 16px; font-weight: 700; font-family: inherit; cursor: pointer; margin-top: 8px; }
.btn:active { background: #27ae60; }
.danger { color: #e74c3c; }
#debug { color: red; font-size: 12px; padding: 8px; }
#dynamic-fields .field { margin-bottom: 16px; }
</style>
</head><body>
<div class="c">
<h2>สร้างเอกสาร</h2>
<p class="sub">เลือกประเภทเอกสารและกรอกข้อมูล</p>
<div id="debug"></div>
<div>
<label>ประเภทเอกสาร <span class="danger">*</span></label>
<select id="sel-category" onchange="onCatChange()">
<option value="">-- เลือกประเภท --</option>
<option value="contract">สัญญา</option>
<option value="letter">หนังสือ</option>
<option value="petition">คำร้อง/คำฟ้อง</option>
<option value="will">พินัยกรรม</option>
<option value="court_form">แบบพิมพ์ศาล</option>
</select>
</div>
<div id="tmpl-wrap" style="display:none;">
<label>รูปแบบเอกสาร <span class="danger">*</span></label>
<select id="sel-tmpl" onchange="onTmplChange()"><option value="">-- เลือกรูปแบบ --</option></select>
<p id="tmpl-desc" style="font-size:13px;color:#7f8c8d;"></p>
</div>
<div id="dynamic-fields"></div>
<div id="submit-wrap" style="display:none;">
<button class="btn" onclick="doSubmit()">สร้างเอกสาร</button>
</div>
</div>
<script>
var T=[], cur=null;
function dbg(m){document.getElementById('debug').textContent+=m+' | ';}
function esc(s){if(!s)return '';var d=document.createElement('div');d.textContent=s;return d.innerHTML;}

dbg('loading...');
fetch('/liff/document/create/data',{method:'POST',headers:{'Content-Type':'application/json'},body:JSON.stringify({jsonrpc:'2.0',method:'call',params:{line_user_id:''}})})
.then(function(r){return r.json();})
.then(function(j){
  dbg('API ok');
  if(j.error){dbg('err:'+j.error.message);return;}
  T=j.result.templates||[];
  dbg(T.length+' templates');
})
.catch(function(e){dbg('fetch err:'+e.message);});

function onCatChange(){
  var c=document.getElementById('sel-category').value;
  var w=document.getElementById('tmpl-wrap');
  var s=document.getElementById('sel-tmpl');
  document.getElementById('dynamic-fields').innerHTML='';
  document.getElementById('submit-wrap').style.display='none';
  cur=null;
  if(!c){w.style.display='none';return;}
  var f=T.filter(function(t){return t.category===c;});
  s.innerHTML='<option value="">-- เลือกรูปแบบ --</option>';
  f.forEach(function(t){s.innerHTML+='<option value="'+t.id+'">'+esc(t.name)+'</option>';});
  w.style.display='block';
}
function onTmplChange(){
  var id=parseInt(document.getElementById('sel-tmpl').value);
  var con=document.getElementById('dynamic-fields');
  con.innerHTML='';cur=null;
  document.getElementById('submit-wrap').style.display='none';
  if(!id)return;
  cur=T.find(function(t){return t.id===id;});
  if(!cur)return;
  document.getElementById('tmpl-desc').textContent=cur.description||'';
  var fields=cur.required_fields||[];
  var h='';
  fields.forEach(function(f){
    h+='<div class="field"><label>'+esc(f.label);
    if(f.required)h+=' <span class="danger">*</span>';
    h+='</label>';
    var r=f.required?' required':'';
    if(f.type==='textarea')h+='<textarea name="'+f.name+'" rows="3"'+r+'></textarea>';
    else if(f.type==='date')h+='<input type="date" name="'+f.name+'"'+r+'/>';
    else if(f.type==='number')h+='<input type="number" name="'+f.name+'"'+r+'/>';
    else h+='<input type="text" name="'+f.name+'"'+r+'/>';
    h+='</div>';
  });
  con.innerHTML=h;
  document.getElementById('submit-wrap').style.display='block';
}
function doSubmit(){
  if(!cur)return;
  var vals={};
  (cur.required_fields||[]).forEach(function(f){
    var el=document.querySelector('[name="'+f.name+'"]');
    vals[f.name]=el?el.value:'';
  });
  dbg('submitting...');
  fetch('/liff/document/create/submit',{method:'POST',headers:{'Content-Type':'application/json'},body:JSON.stringify({jsonrpc:'2.0',method:'call',params:{line_user_id:'',template_id:cur.id,field_values:vals,lead_id:null}})})
  .then(function(r){return r.json();})
  .then(function(j){
    if(j.result&&j.result.draft_id){window.location.href='/liff/document/draft/'+j.result.draft_id;}
    else{dbg('error:'+JSON.stringify(j));}
  })
  .catch(function(e){dbg('submit err:'+e.message);});
}
</script>
</body></html>"""
        return request.make_response(html, headers=[('Content-Type', 'text/html')])


    @http.route('/liff/document/create/data', type='json', auth='public', methods=['POST'], csrf=False)
    def liff_document_create_data(self, **kwargs):
        line_user_id = kwargs.get('line_user_id', '')
        partner, role = self._get_partner_from_line(line_user_id)

        # Allow access if: lawyer role OR no line_user_id (direct browser access)
        if line_user_id and (not partner or role != 'lawyer'):
            return {'error': 'unauthorized', 'message': 'เฉพาะทนายความเท่านั้น'}

        templates = request.env['legal.document.template'].sudo().search([('active', '=', True)])
        tmpl_list = []
        for t in templates:
            tmpl_list.append({
                'id': t.id,
                'name': t.name,
                'code': t.code,
                'category': t.category,
                'description': t.description or '',
                'required_fields': json.loads(t.required_fields or '[]'),
                'is_docx_template': bool(t.template_file_path and t.template_file_path.endswith('.docx')),
                'steps': self._get_template_steps(t),
            })

        # Lawyer's cases for optional linking
        user = None
        if partner:
            user = request.env['res.users'].sudo().search(
                [('partner_id', '=', partner.id)], limit=1)
        cases = []
        if user:
            leads = request.env['crm.lead'].sudo().search([
                '|',
                ('user_id', '=', user.id),
                ('collaborator_ids', 'in', user.id),
                ('case_status', 'not in', ['closed_won', 'closed_lost', 'closed_other']),
            ], order='write_date desc', limit=50)
            for lead in leads:
                cases.append({
                    'id': lead.id,
                    'name': lead.name,
                    'partner_name': lead.partner_id.name or '-',
                })

        return {
            'templates': tmpl_list,
            'cases': cases,
            'lawyer_name': partner.name if partner else '',
        }

    @http.route('/liff/document/create/submit', type='json', auth='public', methods=['POST'], csrf=False)
    def liff_document_create_submit(self, **kwargs):
        line_user_id = kwargs.get('line_user_id', '')
        template_id = kwargs.get('template_id')
        field_values = kwargs.get('field_values', {})
        lead_id = kwargs.get('lead_id')

        partner, role = self._get_partner_from_line(line_user_id)
        if line_user_id and (not partner or role != 'lawyer'):
            return {'error': 'unauthorized', 'message': 'เฉพาะทนายความเท่านั้น'}

        # For direct browser access without LINE, use a default partner
        if not partner:
            partner = request.env['res.partner'].sudo().search([], limit=1)

        template = request.env['legal.document.template'].sudo().browse(template_id)
        if not template.exists():
            return {'error': 'not_found', 'message': 'ไม่พบ template'}

        # Create draft record
        draft_name = f"{template.name} - {field_values.get(json.loads(template.required_fields or '[]')[0]['name'], '')} ({fields.Date.today()})"
        draft = request.env['legal.document.draft'].sudo().create({
            'name': draft_name[:128],
            'template_id': template.id,
            'lead_id': lead_id or False,
            'lawyer_partner_id': partner.id,
            'state': 'generating',
            'field_values': json.dumps(field_values, ensure_ascii=False),
        })

        # Create timeline event if linked to a case
        if lead_id:
            request.env['legal.case.event'].sudo().create({
                'lead_id': lead_id,
                'event_type': 'document',
                'title': f'สร้างเอกสาร: {template.name}',
                'description': f'ทนาย {partner.name} สร้างเอกสาร {template.name}',
            })

        # Kick off background generation
        db_name = request.env.cr.dbname
        draft_id = draft.id
        template_code = template.code
        template_name = template.name
        template_file_path = template.template_file_path
        required_fields_json = template.required_fields
        partner_line_user_id = partner.line_user_id or ''

        # Commit so background thread can see the draft record
        request.env.cr.commit()

        thread = threading.Thread(
            target=self._background_generate,
            args=(db_name, draft_id, template_code, template_name,
                  template_file_path, required_fields_json,
                  field_values, lead_id, partner_line_user_id),
            daemon=True,
        )
        thread.start()

        # Return immediately — user sees "กำลังสร้าง" page
        return {
            'draft_id': draft.id,
            'state': 'generating',
        }

    @staticmethod
    def _background_generate(db_name, draft_id, template_code, template_name,
                             template_file_path, required_fields_json,
                             field_values, lead_id, line_user_id):
        """Generate document in background thread, then push LINE message."""
        import traceback
        import odoo
        from odoo.modules.registry import Registry
        _logger.info('[bg-gen] Starting background generate for draft %s (%s)', draft_id, template_code)
        try:
            registry = Registry(db_name)
            with registry.cursor() as cr:
                env = odoo.api.Environment(cr, odoo.SUPERUSER_ID, {})

                draft = env['legal.document.draft'].browse(draft_id)
                if not draft.exists():
                    _logger.warning('[bg-gen] Draft %s not found, aborting', draft_id)
                    return

                # Generate content
                try:
                    _logger.info('[bg-gen] Calling _generate_document_static for draft %s', draft_id)
                    import base64
                    result = LiffController._generate_document_static(
                        env, template_code, template_name,
                        template_file_path, field_values)

                    # Handle DOCX template-fill result (tuple)
                    if isinstance(result, tuple) and result[0] == '__DOCX__':
                        docx_bytes = result[1]
                        safe_name = re.sub(r'[^\w\s\-]', '', template_name).strip()
                        draft.write({
                            'draft_content': f'[เอกสาร DOCX] {template_name} — สร้างจาก template สำเร็จ',
                            'docx_file': base64.b64encode(docx_bytes),
                            'docx_filename': f'{safe_name}.docx',
                            'state': 'draft',
                        })
                        _logger.info('[bg-gen] DOCX template filled: %d bytes for draft %s',
                                     len(docx_bytes), draft_id)
                    else:
                        # Original markdown content
                        content = result or ''
                        _logger.info('[bg-gen] Generated %d chars for draft %s', len(content), draft_id)
                        draft.write({
                            'draft_content': content,
                            'state': 'draft',
                        })

                    # Save initial version (v1)
                    draft._save_version('auto_generated', 'AI สร้างเอกสารอัตโนมัติ')
                except Exception as e:
                    _logger.error('[bg-gen] Generation error for draft %s: %s\n%s', draft_id, e, traceback.format_exc())
                    draft.write({
                        'draft_content': f'เกิดข้อผิดพลาดในการสร้างเอกสาร: {e}',
                        'state': 'draft',
                    })

                cr.commit()
                _logger.info('[bg-gen] Draft %s committed successfully', draft_id)

                # Push LINE message with link
                if line_user_id:
                    LiffController._push_document_ready(
                        env, line_user_id, draft_id, template_name)

        except Exception as e:
            _logger.error('[bg-gen] Thread error for draft %s: %s\n%s', draft_id, e, traceback.format_exc())

    @staticmethod
    def _push_document_ready(env, line_user_id, draft_id, template_name):
        """Send LINE push notification that document is ready."""
        try:
            import requests as req
            access_token = env['ir.config_parameter'].get_param(
                'line_integration.channel_access_token', '')
            base_url = env['ir.config_parameter'].get_param(
                'web.base.url', 'https://legal.eformservice.com')
            if not access_token:
                return

            doc_url = f'{base_url}/liff/document/draft/{draft_id}'
            resp = req.post(
                'https://api.line.me/v2/bot/message/push',
                headers={
                    'Authorization': f'Bearer {access_token}',
                    'Content-Type': 'application/json',
                },
                json={
                    'to': line_user_id,
                    'messages': [{
                        'type': 'flex',
                        'altText': f'เอกสาร {template_name} พร้อมแล้ว',
                        'contents': {
                            'type': 'bubble',
                            'size': 'kilo',
                            'header': {
                                'type': 'box',
                                'layout': 'vertical',
                                'contents': [{
                                    'type': 'text',
                                    'text': 'เอกสารพร้อมแล้ว',
                                    'weight': 'bold',
                                    'size': 'lg',
                                    'color': '#2ECC71',
                                }],
                            },
                            'body': {
                                'type': 'box',
                                'layout': 'vertical',
                                'contents': [
                                    {
                                        'type': 'text',
                                        'text': template_name,
                                        'weight': 'bold',
                                        'size': 'md',
                                        'wrap': True,
                                    },
                                    {
                                        'type': 'text',
                                        'text': 'ระบบสร้างเอกสารเรียบร้อยแล้ว กดดูเอกสารด้านล่าง',
                                        'size': 'sm',
                                        'color': '#7f8c8d',
                                        'wrap': True,
                                        'margin': 'md',
                                    },
                                ],
                            },
                            'footer': {
                                'type': 'box',
                                'layout': 'vertical',
                                'contents': [{
                                    'type': 'button',
                                    'action': {
                                        'type': 'uri',
                                        'label': 'ดูเอกสาร',
                                        'uri': doc_url,
                                    },
                                    'style': 'primary',
                                    'color': '#2ECC71',
                                }],
                            },
                        },
                    }],
                },
                timeout=10,
            )
            if resp.status_code != 200:
                _logger.error('LINE push document ready failed: %s', resp.text)
        except Exception as e:
            _logger.error('LINE push document ready error: %s', e)

    @staticmethod
    def _fill_docx_template(template_path, field_values):
        """Fill a .docx master template: replace {{key}} with values. Returns bytes."""
        from docxtpl import DocxTemplate
        from ..utils.thai_number import number_to_thai_text, baht_text

        # Auto-generate _text (ตัวหนังสือ) and _baht (บาท) for numeric values
        enriched = dict(field_values)
        for key, val in field_values.items():
            try:
                num = float(str(val).replace(',', ''))
                enriched[key + '_text'] = number_to_thai_text(num)
                enriched[key + '_baht'] = baht_text(num)
            except (ValueError, TypeError):
                pass

        # Auto-expand address fields (JSON from frontend: {"house":"..","t":"..","a":"..","p":"..","z":".."})
        for key, val in list(field_values.items()):
            if isinstance(val, str) and val.startswith('{') and '"t"' in val:
                try:
                    addr = json.loads(val)
                    if 't' in addr and 'a' in addr:
                        house = addr.get('house', '')
                        tambon = addr.get('t', '')
                        amphoe = addr.get('a', '')
                        province = addr.get('p', '')
                        zipcode = addr.get('z', '')
                        enriched[key + '_เลขที่'] = house
                        enriched[key + '_ตำบล'] = tambon
                        enriched[key + '_อำเภอ'] = amphoe
                        enriched[key + '_จังหวัด'] = province
                        enriched[key + '_รหัสไปรษณีย์'] = zipcode
                        # Full address string
                        parts = [p for p in [
                            house,
                            f'ตำบล{tambon}' if tambon else '',
                            f'อำเภอ{amphoe}' if amphoe else '',
                            f'จังหวัด{province}' if province else '',
                            zipcode,
                        ] if p]
                        enriched[key] = ' '.join(parts)
                except (json.JSONDecodeError, TypeError):
                    pass

        doc = DocxTemplate(template_path)
        # Register Jinja2 filters for use in templates: {{amount|baht_text}}
        doc.jinja_env.filters['number_text'] = number_to_thai_text
        doc.jinja_env.filters['baht_text'] = baht_text
        doc.render(enriched)

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    @staticmethod
    def _render_md_jinja2(content, field_values):
        """Render markdown template with Jinja2 (conditional logic + loops).

        Supports:
        - {% if condition %} ... {% endif %}
        - {% for item in list %} ... {% endfor %}
        - {{variable}} placeholders
        - {% set clause_no = namespace(n=1) %} for auto-numbering
        - Filters: |number_text, |baht_text
        """
        from jinja2 import Environment, BaseLoader, Undefined

        class SilentUndefined(Undefined):
            """Return empty string for undefined variables instead of error."""
            def __str__(self):
                return ''
            def __iter__(self):
                return iter([])
            def __bool__(self):
                return False

        env = Environment(
            loader=BaseLoader(),
            undefined=SilentUndefined,
            keep_trailing_newline=True,
        )

        # Register Thai number filters
        try:
            from ..utils.thai_number import number_to_thai_text, baht_text
            env.filters['number_text'] = number_to_thai_text
            env.filters['baht_text'] = baht_text
        except ImportError:
            pass

        # Enrich field_values: auto-generate _text and _baht for numbers
        enriched = dict(field_values)
        for key, val in field_values.items():
            try:
                num = float(str(val).replace(',', ''))
                try:
                    from ..utils.thai_number import number_to_thai_text, baht_text
                    enriched[key + '_text'] = number_to_thai_text(num)
                    enriched[key + '_baht'] = baht_text(num)
                except ImportError:
                    pass
            except (ValueError, TypeError):
                pass

            # Parse JSON arrays for repeating sections
            if isinstance(val, str):
                val_stripped = val.strip()
                if val_stripped.startswith('[') and val_stripped.endswith(']'):
                    try:
                        parsed = json.loads(val_stripped)
                        if isinstance(parsed, list):
                            enriched[key] = parsed
                    except (json.JSONDecodeError, TypeError):
                        pass
                elif val_stripped.startswith('{') and val_stripped.endswith('}'):
                    try:
                        parsed = json.loads(val_stripped)
                        if isinstance(parsed, dict):
                            enriched[key] = parsed
                    except (json.JSONDecodeError, TypeError):
                        pass

        try:
            template = env.from_string(content)
            rendered = template.render(**enriched)
            # Clean up excessive blank lines from conditional blocks
            rendered = re.sub(r'\n{3,}', '\n\n', rendered)
            return rendered
        except Exception as e:
            _logger.warning('Jinja2 render failed, falling back to simple replace: %s', e)
            # Fallback: simple string replace
            for key, val in field_values.items():
                content = content.replace('{{' + key + '}}', str(val))
            return content

    @staticmethod
    def _generate_document_static(env, template_code, template_name,
                                  template_file_path, field_values):
        """Generate document content (static method for background thread).

        If template_file_path is .docx → template-fill (fast, no AI cost).
        Returns either a string (markdown) or tuple ('__DOCX__', bytes).
        """
        base_path = '/mnt/templates'

        # ── NEW: DOCX template-fill approach ──
        if template_file_path and template_file_path.endswith('.docx'):
            tmpl_path = os.path.join(base_path, template_file_path)
            if os.path.exists(tmpl_path):
                try:
                    docx_bytes = LiffController._fill_docx_template(
                        tmpl_path, field_values)
                    if docx_bytes:
                        _logger.info('[docx-fill] Generated %d bytes for %s',
                                     len(docx_bytes), template_code)
                        return ('__DOCX__', docx_bytes)
                except Exception as e:
                    _logger.warning('[docx-fill] Failed for %s: %s', template_code, e)

        # ── ORIGINAL: Try adkcode API ──
        adkcode_url = env['ir.config_parameter'].get_param(
            'line_integration.adkcode_url', '')
        if adkcode_url:
            try:
                import requests as req
                resp = req.post(
                    f'{adkcode_url}/draft-document',
                    json={
                        'template_code': template_code,
                        'fields': field_values,
                    },
                    timeout=120,
                )
                if resp.status_code == 200:
                    data = resp.json()
                    if data.get('content'):
                        return data['content']
            except Exception as e:
                _logger.warning('adkcode draft-document failed, falling back: %s', e)

        # Fallback: read .md template file and fill with Jinja2
        tmpl_path = os.path.join(base_path, template_file_path or template_code + '.md')

        if os.path.exists(tmpl_path):
            with open(tmpl_path, 'r', encoding='utf-8') as f:
                content = f.read()
            content = re.sub(r'^---\n.*?\n---\n', '', content, flags=re.DOTALL)
            content = LiffController._render_md_jinja2(content, field_values)
            return content.strip()

        # No template file — generate simple document
        lines = [f'# {template_name}\n']
        for key, val in field_values.items():
            lines.append(f'**{key}:** {val}')
        return '\n'.join(lines)

    # ── /liff/document/draft/<id>/status (poll for generation) ──

    @http.route('/liff/document/draft/<int:draft_id>/status', type='json', auth='public', methods=['POST'], csrf=False)
    def liff_document_draft_status(self, draft_id, **kwargs):
        """Poll endpoint: returns draft state for JS polling."""
        draft = request.env['legal.document.draft'].sudo().browse(draft_id)
        if not draft.exists():
            return {'state': 'not_found'}
        return {
            'state': draft.state,
            'content_length': len(draft.draft_content or ''),
            'has_docx': bool(draft.docx_file),
            'template_name': draft.template_id.name or '',
            'docx_filename': draft.docx_filename or '',
        }

    # ── /liff/document/draft/<id> ─────────────────────────────

    @http.route('/liff/document/draft/<int:draft_id>', type='http', auth='public', methods=['GET'], csrf=False)
    def liff_document_draft(self, draft_id, **kwargs):
        draft = request.env['legal.document.draft'].sudo().browse(draft_id)
        if not draft.exists():
            return request.render('legal_liff.liff_error_page', {
                'message': 'ไม่พบเอกสาร',
                'liff_id': self._get_liff_id(),
            })
        return request.render('legal_liff.liff_document_draft_page', {
            'draft': draft,
            'liff_id': self._get_liff_id(),
        })

    # ── /liff/document/draft/<id>/download ────────────────────

    @http.route('/liff/document/draft/<int:draft_id>/download',
                type='http', auth='public', methods=['GET'], csrf=False)
    def liff_document_download(self, draft_id, **kwargs):
        """Download document draft as PDF or DOCX."""
        import base64
        draft = request.env['legal.document.draft'].sudo().browse(draft_id)
        if not draft.exists():
            return request.not_found()

        # If draft has a stored DOCX file (template-based), use it directly
        has_docx = bool(draft.docx_file)

        if not has_docx and not draft.draft_content:
            return request.not_found()

        fmt = (kwargs.get('format') or 'pdf').lower()
        if fmt not in ('pdf', 'docx'):
            fmt = 'pdf'

        if has_docx:
            docx_bytes = base64.b64decode(draft.docx_file)
            if fmt == 'docx':
                data = docx_bytes
                content_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                ext = 'docx'
            else:
                # Convert DOCX → PDF via libreoffice
                data, content_type, ext = self._convert_docx_to_pdf(docx_bytes)
        elif fmt == 'pdf':
            data, content_type, ext = self._render_pdf(draft)
        else:
            data, content_type, ext = self._render_docx(draft)

        # Use ASCII-safe filename + UTF-8 filename* for Thai support
        ascii_name = f'document_{draft_id}.{ext}'
        from urllib.parse import quote
        utf8_name = quote(f'{draft.name or "document"}.{ext}')
        disposition = f"attachment; filename=\"{ascii_name}\"; filename*=UTF-8''{utf8_name}"

        return request.make_response(data, headers=[
            ('Content-Type', content_type),
            ('Content-Disposition', disposition),
            ('Content-Length', str(len(data))),
        ])

    def _convert_docx_to_pdf(self, docx_bytes):
        """Convert DOCX bytes to PDF: read docx → HTML → wkhtmltopdf."""
        from docx import Document as DocxDoc

        doc = DocxDoc(io.BytesIO(docx_bytes))

        # Build HTML from docx paragraphs
        html_parts = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                html_parts.append('<p>&nbsp;</p>')
                continue

            # Detect alignment
            align = ''
            if para.alignment == 1:  # CENTER
                align = 'text-align:center;'
            elif para.alignment == 2:  # RIGHT
                align = 'text-align:right;'

            # Detect bold title
            is_bold = all(r.bold for r in para.runs if r.text.strip()) if para.runs else False
            tag = 'p'
            extra_style = ''
            if is_bold and len(text) < 40:
                extra_style = 'font-weight:bold; font-size:20px;'

            # Detect first-line indent
            indent = ''
            pf = para.paragraph_format
            if pf.first_line_indent and pf.first_line_indent > 0:
                indent = 'text-indent:1.5cm;'

            # Check for underline in runs
            parts = []
            for run in para.runs:
                t = run.text
                if not t:
                    continue
                if run.font.underline:
                    parts.append(f'<u>{t}</u>')
                elif run.font.bold:
                    parts.append(f'<b>{t}</b>')
                else:
                    parts.append(t)

            inner = ''.join(parts) if parts else text
            style = f'{align}{indent}{extra_style}'.strip()
            style_attr = f' style="{style}"' if style else ''
            html_parts.append(f'<{tag}{style_attr}>{inner}</{tag}>')

        html_body = '\n'.join(html_parts)
        html = f"""<!DOCTYPE html>
<html><head><meta charset="utf-8"/>
<style>
@import url('https://fonts.googleapis.com/css2?family=Noto+Sans+Thai:wght@400;700&display=swap');
body {{ font-family: 'Noto Sans Thai', 'TH Sarabun New', sans-serif; font-size: 16px; line-height: 1.6; margin: 40px; color: #333; }}
p {{ margin: 2px 0; }}
u {{ text-decoration: underline; }}
</style></head>
<body>{html_body}</body></html>"""

        with tempfile.NamedTemporaryFile(suffix='.html', delete=False) as f_in:
            f_in.write(html.encode('utf-8'))
            html_path = f_in.name
        pdf_path = html_path.replace('.html', '.pdf')

        try:
            subprocess.run([
                '/usr/local/bin/wkhtmltopdf', '--quiet',
                '--encoding', 'utf-8',
                '--page-size', 'A4',
                '--margin-top', '15mm', '--margin-bottom', '15mm',
                '--margin-left', '20mm', '--margin-right', '15mm',
                html_path, pdf_path,
            ], check=True, timeout=30)

            with open(pdf_path, 'rb') as f:
                pdf_data = f.read()
        finally:
            for p in (html_path, pdf_path):
                try:
                    os.unlink(p)
                except OSError:
                    pass

        return pdf_data, 'application/pdf', 'pdf'

    def _render_pdf(self, draft):
        """Render draft content to PDF using wkhtmltopdf."""
        import markdown as md

        html_body = md.markdown(
            draft.draft_content or '',
            extensions=['tables', 'nl2br'],
        )

        html = f"""<!DOCTYPE html>
<html><head><meta charset="utf-8"/>
<style>
@import url('https://fonts.googleapis.com/css2?family=Noto+Sans+Thai:wght@400;700&display=swap');
body {{ font-family: 'Noto Sans Thai', 'TH Sarabun New', sans-serif; font-size: 14px; line-height: 1.8; margin: 40px; color: #333; }}
h1 {{ font-size: 22px; text-align: center; margin-bottom: 24px; }}
h2 {{ font-size: 18px; margin-top: 20px; }}
table {{ border-collapse: collapse; width: 100%; margin: 12px 0; }}
th, td {{ border: 1px solid #ccc; padding: 8px; text-align: left; }}
th {{ background: #f5f5f5; }}
</style></head>
<body>
<h1>{draft.template_id.name or draft.name or 'เอกสาร'}</h1>
{html_body}
</body></html>"""

        with tempfile.NamedTemporaryFile(suffix='.html', delete=False) as f_in:
            f_in.write(html.encode('utf-8'))
            html_path = f_in.name
        pdf_path = html_path.replace('.html', '.pdf')

        try:
            subprocess.run([
                '/usr/local/bin/wkhtmltopdf',
                '--quiet',
                '--encoding', 'utf-8',
                '--page-size', 'A4',
                '--margin-top', '15mm',
                '--margin-bottom', '15mm',
                '--margin-left', '15mm',
                '--margin-right', '15mm',
                html_path, pdf_path,
            ], check=True, timeout=30)

            with open(pdf_path, 'rb') as f:
                pdf_data = f.read()
        finally:
            for p in (html_path, pdf_path):
                try:
                    os.unlink(p)
                except OSError:
                    pass

        return pdf_data, 'application/pdf', 'pdf'

    def _render_docx(self, draft):
        """Render draft content to DOCX."""
        from docx import Document as DocxDocument
        from docx.shared import Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = DocxDocument()

        # Page margins
        for section in doc.sections:
            section.top_margin = Cm(2.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)

        # Title
        title = doc.add_heading(
            draft.template_id.name or draft.name or 'เอกสาร', level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Content — parse line by line
        content = draft.draft_content or ''
        for line in content.split('\n'):
            stripped = line.strip()
            if not stripped:
                doc.add_paragraph('')
                continue
            if stripped.startswith('## '):
                doc.add_heading(stripped[3:], level=2)
            elif stripped.startswith('# '):
                doc.add_heading(stripped[2:], level=1)
            elif stripped.startswith('- ') or stripped.startswith('* '):
                doc.add_paragraph(stripped[2:], style='List Bullet')
            else:
                p = doc.add_paragraph(stripped)
                for run in p.runs:
                    run.font.size = Pt(14)

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue(), \
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document', \
            'docx'

    # ── /liff/document/draft/<id>/action ──────────────────────

    @http.route('/liff/document/draft/<int:draft_id>/action',
                type='json', auth='public', methods=['POST'], csrf=False)
    def liff_document_action(self, draft_id, **kwargs):
        """Handle document state transitions."""
        action = kwargs.get('action', '')
        notes = kwargs.get('notes', '')

        draft = request.env['legal.document.draft'].sudo().browse(draft_id)
        if not draft.exists():
            return {'success': False, 'error': 'ไม่พบเอกสาร'}

        action_map = {
            'send_to_client': draft.action_send_to_client,
            'request_revision': lambda: draft.action_request_revision(notes),
            'finalize': draft.action_finalize,
            'sign': draft.action_sign,
            'cancel': draft.action_cancel,
            'back_to_draft': draft.action_back_to_draft,
        }

        handler = action_map.get(action)
        if not handler:
            return {'success': False, 'error': f'ไม่รู้จัก action: {action}'}

        result = handler()
        if not result:
            return {'success': False, 'error': f'ไม่สามารถ {action} จากสถานะ {draft.state}'}

        return {
            'success': True,
            'new_state': draft.state,
            'message': f'เปลี่ยนสถานะเป็น {dict(draft._fields["state"].selection).get(draft.state, draft.state)}',
        }

    # ── /api/documents (REST JSON for LINE bot) ────────────────

    @http.route('/api/documents', type='http', auth='public', methods=['GET'], csrf=False)
    def api_document_list(self, line_user_id='', **kwargs):
        """REST API: list documents. Called by LINE bot."""
        base_url = request.env['ir.config_parameter'].sudo().get_param(
            'web.base.url', 'https://legal.eformservice.com')
        # Ensure HTTPS for external URLs
        base_url = base_url.replace('http://', 'https://')

        if line_user_id:
            partner = request.env['res.partner'].sudo().search(
                [('line_user_id', '=', line_user_id)], limit=1)
        else:
            partner = None

        # If partner found, filter by role; otherwise show recent drafts
        if partner and partner.line_role == 'lawyer':
            # Lawyer sees all non-generating documents
            domain = [('state', '!=', 'generating')]
        elif partner:
            domain = [
                ('client_partner_id', '=', partner.id),
                ('state', 'in', ('sent_to_client', 'revision', 'final', 'signed', 'draft')),
            ]
        else:
            domain = [('state', '!=', 'generating')]

        docs = request.env['legal.document.draft'].sudo().search(
            domain, order='create_date desc', limit=20)

        state_labels = {
            'generating': 'กำลังสร้าง', 'draft': 'แบบร่าง',
            'sent_to_client': 'ส่งให้ลูกค้าแล้ว', 'revision': 'แก้ไข',
            'final': 'ฉบับสมบูรณ์', 'signed': 'เซ็นแล้ว',
            'cancelled': 'ยกเลิก', 'expired': 'หมดอายุ',
        }

        result = {
            'total': len(docs),
            'base_url': base_url,
            'documents': [{
                'id': d.id,
                'name': d.name or '',
                'state': d.state,
                'state_label': state_labels.get(d.state, d.state),
                'template_name': d.template_id.name or '',
                'create_date': str(d.create_date)[:16] if d.create_date else '',
                'view_url': f'{base_url}/liff/document/draft/{d.id}',
                'pdf_url': f'{base_url}/liff/document/draft/{d.id}/download?fmt=pdf',
                'content_length': len(d.draft_content or ''),
            } for d in docs],
        }
        return request.make_json_response(result)

    # ── /liff/documents (document list per user) ──────────────

    @http.route('/liff/documents', type='http', auth='public', methods=['GET'], csrf=False)
    def liff_document_list(self, **kwargs):
        """List all documents for a user."""
        return request.render('legal_liff.liff_document_list_page', {
            'liff_id': self._get_liff_id(),
        })

    @http.route('/liff/documents/data', type='json', auth='public', methods=['POST'], csrf=False)
    def liff_document_list_data(self, **kwargs):
        """JSON: get documents for a LINE user."""
        line_user_id = kwargs.get('line_user_id', '')
        if not line_user_id:
            return {'documents': []}

        partner = request.env['res.partner'].sudo().search(
            [('line_user_id', '=', line_user_id)], limit=1)
        if not partner:
            return {'documents': []}

        # Lawyer sees docs they created; client sees docs sent to them
        if partner.line_role == 'lawyer':
            domain = [('lawyer_partner_id', '=', partner.id)]
        else:
            domain = [
                ('client_partner_id', '=', partner.id),
                ('state', 'in', ('sent_to_client', 'revision', 'final', 'signed')),
            ]

        docs = request.env['legal.document.draft'].sudo().search(
            domain, order='create_date desc', limit=50)

        state_labels = dict(docs._fields['state'].selection)
        return {
            'documents': [{
                'id': d.id,
                'name': d.name,
                'state': d.state,
                'state_label': state_labels.get(d.state, d.state),
                'template_name': d.template_id.name or '',
                'create_date': str(d.create_date)[:16],
                'revision_count': d.revision_count,
            } for d in docs],
        }

    # ── Address Search ────────────────────────────────────────

    _address_data = None  # class-level cache

    @classmethod
    def _load_address_data(cls):
        """Load Thai address database (cached)."""
        if cls._address_data is not None:
            return cls._address_data
        data_path = os.path.join(
            os.path.dirname(os.path.dirname(__file__)),
            'data', 'thai_address.json',
        )
        try:
            with open(data_path, 'r', encoding='utf-8') as f:
                cls._address_data = json.load(f)
        except (FileNotFoundError, json.JSONDecodeError) as e:
            _logger.warning('Thai address data not found: %s', e)
            cls._address_data = []
        return cls._address_data

    @http.route('/liff/address/search', type='json', auth='public', csrf=False)
    def liff_address_search(self, q='', limit=20, **kw):
        """Search Thai address by tambon/amphoe/province/zip."""
        if not q or len(q) < 2:
            return {'results': []}

        data = self._load_address_data()
        q_lower = q.lower().strip()
        results = []
        for addr in data:
            if (q_lower in addr.get('t', '').lower()
                    or q_lower in addr.get('a', '').lower()
                    or q_lower in addr.get('p', '').lower()
                    or q_lower in str(addr.get('z', ''))):
                results.append(addr)
                if len(results) >= limit:
                    break
        return {'results': results}

    # ═══ CLAUSE SUGGESTION API ═══════════════════════════════

    @http.route('/liff/clause/suggest', type='json', auth='public',
                methods=['POST'], csrf=False)
    def liff_clause_suggest(self, **kwargs):
        """Suggest clauses from the clause library for a given context.

        Input:
            document_type: 'contract' | 'letter' | 'petition' | 'will'
            category_code: optional, e.g. 'guarantee', 'penalty'
            query: optional search text
            risk_level: optional 'conservative' | 'standard' | 'aggressive'
            template_id: optional, filter by compatible template

        Returns:
            clauses: [{id, name, code, content, risk_level, category, legal_reference}]
        """
        doc_type = kwargs.get('document_type', 'all')
        category_code = kwargs.get('category_code', '')
        query = kwargs.get('query', '')
        risk_level = kwargs.get('risk_level', '')
        template_id = kwargs.get('template_id')

        Clause = request.env['legal.clause'].sudo()

        domain = [('state', '=', 'approved')]
        if doc_type and doc_type != 'all':
            domain.append(('document_types', 'in', [doc_type, 'all']))
        if category_code:
            domain.append(('category_id.code', '=', category_code))
        if risk_level:
            domain.append(('risk_level', '=', risk_level))
        if template_id:
            domain.append(('template_ids', 'in', [template_id]))

        clauses = Clause.search(domain, order='usage_count desc, sequence', limit=20)

        # Filter by search query (name, tags, content)
        if query:
            q = query.lower()
            clauses = clauses.filtered(
                lambda c: q in (c.name or '').lower()
                or q in (c.tags or '').lower()
                or q in (c.content or '').lower()
            )

        result = []
        for c in clauses:
            result.append({
                'id': c.id,
                'name': c.name,
                'code': c.code,
                'content': c.content,
                'risk_level': c.risk_level,
                'category': c.category_id.name if c.category_id else '',
                'category_code': c.category_id.code if c.category_id else '',
                'legal_reference': c.legal_reference or '',
                'usage_count': c.usage_count,
                'tags': c.tags or '',
            })

        return {'clauses': result}

    @http.route('/liff/clause/use', type='json', auth='public',
                methods=['POST'], csrf=False)
    def liff_clause_use(self, **kwargs):
        """Record that a clause was used (increment counter)."""
        clause_id = kwargs.get('clause_id')
        if clause_id:
            clause = request.env['legal.clause'].sudo().browse(clause_id)
            if clause.exists():
                clause.increment_usage()
                return {'success': True}
        return {'success': False}

    @http.route('/liff/clause/ai-suggest', type='json', auth='public',
                methods=['POST'], csrf=False)
    def liff_clause_ai_suggest(self, **kwargs):
        """Use AI to suggest a clause based on context + clause library.

        Input:
            document_type: e.g. 'contract'
            clause_topic: e.g. 'หลักประกัน', 'เบี้ยปรับ'
            context: e.g. 'สัญญาเช่าอาคาร มูลค่า 5 ล้านบาท'
            risk_level: optional preferred risk level

        Returns:
            suggestions: [
                {source: 'library', clause: {...}},
                {source: 'ai', content: '...'},
            ]
        """
        doc_type = kwargs.get('document_type', 'contract')
        clause_topic = kwargs.get('clause_topic', '')
        context = kwargs.get('context', '')
        risk_level = kwargs.get('risk_level', 'standard')

        suggestions = []

        # 1) Search clause library first
        Clause = request.env['legal.clause'].sudo()
        domain = [
            ('state', '=', 'approved'),
            ('document_types', 'in', [doc_type, 'all']),
        ]
        library_clauses = Clause.search(domain, order='usage_count desc')
        if clause_topic:
            q = clause_topic.lower()
            library_clauses = library_clauses.filtered(
                lambda c: q in (c.name or '').lower()
                or q in (c.tags or '').lower()
                or q in (c.category_id.name or '').lower()
            )

        # Return up to 3 from library grouped by risk level
        for level in [risk_level, 'standard', 'conservative', 'aggressive']:
            level_clauses = library_clauses.filtered(lambda c: c.risk_level == level)
            for c in level_clauses[:1]:
                if not any(s.get('clause', {}).get('id') == c.id for s in suggestions):
                    suggestions.append({
                        'source': 'library',
                        'clause': {
                            'id': c.id,
                            'name': c.name,
                            'content': c.content,
                            'risk_level': c.risk_level,
                            'legal_reference': c.legal_reference or '',
                        },
                    })
            if len(suggestions) >= 3:
                break

        # 2) Ask AI to generate a custom clause
        adkcode_url = request.env['ir.config_parameter'].sudo().get_param(
            'line_integration.adkcode_url', '')
        if adkcode_url and clause_topic:
            try:
                import requests as req
                prompt = (
                    f'สร้างข้อสัญญาเรื่อง "{clause_topic}" '
                    f'สำหรับ {doc_type} '
                    f'บริบท: {context}\n'
                    f'ระดับความเข้มงวด: {risk_level}\n'
                    f'ตอบเฉพาะเนื้อหาข้อสัญญา ไม่ต้องมีหัวข้อ '
                    f'ใช้ภาษากฎหมายไทยที่ถูกต้อง อ้างอิงมาตราที่เกี่ยวข้อง'
                )
                resp = req.post(
                    f'{adkcode_url}/chat',
                    json={'message': prompt, 'line_user_id': '__clause_gen__'},
                    timeout=30,
                )
                if resp.status_code == 200:
                    data = resp.json()
                    ai_content = data.get('response', data.get('content', ''))
                    if ai_content:
                        suggestions.append({
                            'source': 'ai',
                            'content': ai_content,
                            'note': 'AI แนะนำ — ต้องให้ทนายตรวจสอบก่อนใช้',
                        })
            except Exception as e:
                _logger.warning('AI clause suggestion failed: %s', e)

        return {'suggestions': suggestions}

    @http.route('/liff/document/draft/<int:draft_id>/versions', type='json',
                auth='public', methods=['POST'], csrf=False)
    def liff_document_versions(self, draft_id, **kwargs):
        """Get version history for a document."""
        draft = request.env['legal.document.draft'].sudo().browse(draft_id)
        if not draft.exists():
            return {'versions': []}

        versions = []
        for v in draft.version_ids.sorted(lambda r: r.version_number, reverse=True):
            versions.append({
                'id': v.id,
                'version': v.version_number,
                'state': v.state_at_save,
                'change_type': v.change_type,
                'change_summary': v.change_summary or '',
                'changed_by': v.changed_by.name if v.changed_by else '',
                'date': v.create_date.isoformat() if v.create_date else '',
                'has_content': bool(v.content),
                'has_docx': bool(v.docx_file),
            })

        return {
            'current_version': draft.current_version,
            'versions': versions,
        }
