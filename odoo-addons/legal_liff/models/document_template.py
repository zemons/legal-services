import io
import json
import logging
import os
import re

from odoo import models, fields, api
from odoo.exceptions import UserError

_logger = logging.getLogger(__name__)


class LegalDocumentTemplate(models.Model):
    _name = 'legal.document.template'
    _description = 'Legal Document Template'
    _order = 'category, sequence, name'

    name = fields.Char('Template Name', required=True)
    code = fields.Char('Code', required=True, help='e.g. contract/rental-agreement')
    category = fields.Selection([
        ('contract', 'สัญญา'),
        ('letter', 'หนังสือ'),
        ('petition', 'คำร้อง/คำฟ้อง'),
        ('will', 'พินัยกรรม'),
        ('court_form', 'แบบพิมพ์ศาล'),
    ], string='Category', required=True, index=True)
    description = fields.Text('Description')
    field_ids = fields.One2many(
        'legal.document.template.field', 'template_id',
        string='Template Fields',
    )
    required_fields = fields.Text(
        'Required Fields (JSON)',
        help='JSON array — auto-synced from Template Fields tab',
    )
    template_file_path = fields.Char('Template File Path', help='Relative to data/templates/')
    active = fields.Boolean(default=True)
    sequence = fields.Integer(default=10)

    # ── Upload & AI Processing ────────────────────────────────
    master_docx = fields.Binary('Master DOCX', attachment=True,
                                help='อัพโหลดไฟล์ .docx ต้นแบบ (หัวสำนักงาน/แบบฟอร์มราชการ)')
    master_docx_filename = fields.Char('Master DOCX Filename')
    processed_docx = fields.Binary('Processed DOCX', attachment=True,
                                   help='ไฟล์ที่ AI แปลงเป็น template แล้ว (มี {{placeholder}})')
    processed_docx_filename = fields.Char('Processed DOCX Filename')
    processing_state = fields.Selection([
        ('none', 'ยังไม่อัพโหลด'),
        ('uploaded', 'อัพโหลดแล้ว'),
        ('processing', 'AI กำลังวิเคราะห์'),
        ('ready', 'พร้อมใช้งาน'),
        ('error', 'เกิดข้อผิดพลาด'),
    ], string='Processing State', default='none', index=True)
    processing_log = fields.Text('Processing Log')

    # ── Computed ──────────────────────────────────────────────
    required_fields_count = fields.Integer(
        'Field Count', compute='_compute_required_fields_count')

    # ── Sync helpers ──────────────────────────────────────────

    def _sync_json_from_fields(self):
        """Rebuild required_fields JSON from field_ids."""
        for rec in self:
            fields_list = []
            for f in rec.field_ids.sorted('sequence'):
                entry = {
                    'name': f.name,
                    'label': f.label,
                    'type': f.field_type,
                    'required': f.required,
                }
                if f.field_type == 'select' and f.options:
                    entry['options'] = [o.strip() for o in f.options.split(',') if o.strip()]
                if f.show_when:
                    try:
                        entry['show_when'] = json.loads(f.show_when)
                    except (json.JSONDecodeError, TypeError):
                        pass
                fields_list.append(entry)
            rec.required_fields = json.dumps(fields_list, ensure_ascii=False, indent=2) if fields_list else '[]'

    def _sync_fields_from_json(self):
        """Create/replace field_ids from required_fields JSON."""
        FieldModel = self.env['legal.document.template.field']
        for rec in self:
            try:
                fields_list = json.loads(rec.required_fields or '[]')
            except (json.JSONDecodeError, TypeError):
                continue
            if not isinstance(fields_list, list):
                continue
            rec.field_ids.unlink()
            for i, f in enumerate(fields_list):
                options_val = ''
                if isinstance(f.get('options'), list):
                    options_val = ', '.join(str(o) for o in f['options'])
                show_when_val = ''
                if f.get('show_when'):
                    show_when_val = json.dumps(f['show_when'], ensure_ascii=False)
                FieldModel.create({
                    'template_id': rec.id,
                    'sequence': (i + 1) * 10,
                    'name': f.get('name', ''),
                    'label': f.get('label', f.get('name', '')),
                    'field_type': f.get('type', 'text'),
                    'required': f.get('required', True),
                    'options': options_val,
                    'show_when': show_when_val,
                })

    @api.model_create_multi
    def create(self, vals_list):
        records = super().create(vals_list)
        # If created with required_fields JSON (e.g. XML data / AI), sync to field_ids
        for rec in records:
            if rec.required_fields and not rec.field_ids:
                rec._sync_fields_from_json()
        return records

    def write(self, vals):
        res = super().write(vals)
        if 'field_ids' in vals:
            # Admin edited the tree → rebuild JSON
            self._sync_json_from_fields()
        elif 'required_fields' in vals and 'field_ids' not in vals:
            # AI or direct JSON edit → rebuild field_ids
            self._sync_fields_from_json()
        return res

    @api.depends('field_ids')
    def _compute_required_fields_count(self):
        for rec in self:
            rec.required_fields_count = len(rec.field_ids)

    @api.onchange('master_docx')
    def _onchange_master_docx(self):
        if self.master_docx:
            self.processing_state = 'uploaded'

    # ── Actions ───────────────────────────────────────────────

    def action_ai_analyze(self):
        """Button action: AI วิเคราะห์ไฟล์ต้นแบบ → สร้าง template + required_fields"""
        self.ensure_one()
        if not self.master_docx:
            raise UserError('กรุณาอัพโหลดไฟล์ .docx ก่อน')

        self.write({
            'processing_state': 'processing',
            'processing_log': 'เริ่มวิเคราะห์...\n',
        })

        try:
            self._run_ai_analysis()
        except Exception as e:
            _logger.error('AI analysis failed for template %s: %s', self.id, e)
            self.write({
                'processing_state': 'error',
                'processing_log': (self.processing_log or '') + f'\nError: {e}',
            })

    def _convert_doc_to_docx(self, doc_bytes):
        """Convert old .doc (OLE2) to .docx using LibreOffice."""
        import subprocess
        import tempfile
        import base64

        with tempfile.TemporaryDirectory() as tmpdir:
            input_path = os.path.join(tmpdir, 'input.doc')
            with open(input_path, 'wb') as f:
                f.write(doc_bytes)

            result = subprocess.run([
                'libreoffice', '--headless', '--convert-to', 'docx',
                '--outdir', tmpdir, input_path,
            ], capture_output=True, timeout=60)

            output_path = os.path.join(tmpdir, 'input.docx')
            if result.returncode != 0 or not os.path.exists(output_path):
                raise UserError(
                    'ไม่สามารถแปลงไฟล์ .doc เป็น .docx ได้\n'
                    f'{result.stderr.decode(errors="replace")}'
                )

            with open(output_path, 'rb') as f:
                return f.read()

    def _run_ai_analysis(self):
        """Read uploaded docx, send text to Gemini, get placeholders & fields."""
        import base64
        from docx import Document as DocxDoc

        # 1) Read docx content — auto-convert .doc to .docx if needed
        docx_bytes = base64.b64decode(self.master_docx)

        # Detect old .doc format (OLE2 magic bytes: D0 CF 11 E0)
        if docx_bytes[:4] == b'\xd0\xcf\x11\xe0':
            self.write({
                'processing_log': (self.processing_log or '') +
                                  'ตรวจพบไฟล์ .doc (Word 97-2003) กำลังแปลงเป็น .docx...\n',
            })
            docx_bytes = self._convert_doc_to_docx(docx_bytes)
            # Update the stored master_docx to .docx format
            self.write({
                'master_docx': base64.b64encode(docx_bytes),
                'master_docx_filename': (self.master_docx_filename or 'file').rsplit('.', 1)[0] + '.docx',
                'processing_log': (self.processing_log or '') + 'แปลงเป็น .docx สำเร็จ\n',
            })

        doc = DocxDoc(io.BytesIO(docx_bytes))

        paragraphs = []
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if text:
                paragraphs.append(f'[{i}] {text}')

        doc_text = '\n'.join(paragraphs)
        self.write({
            'processing_log': (self.processing_log or '') +
                              f'อ่านเอกสารได้ {len(paragraphs)} ย่อหน้า\n',
        })

        # 2) Call Gemini API
        api_key = self.env['ir.config_parameter'].sudo().get_param(
            'legal_liff.google_api_key', '')
        if not api_key:
            # Fallback: try env var from adkcode config
            api_key = os.environ.get('GOOGLE_API_KEY', '')
        if not api_key:
            raise UserError(
                'ไม่พบ Google API Key\n'
                'ตั้งค่าที่: Settings > System Parameters > legal_liff.google_api_key'
            )

        model_name = self.env['ir.config_parameter'].sudo().get_param(
            'legal_liff.gemini_model', 'gemini-2.5-flash')

        prompt = self._build_analysis_prompt(doc_text)

        # 2-3) Call Gemini with retry on JSON parse failure
        max_retries = 3
        last_error = None
        reconstructed = None
        required_fields_list = None

        for attempt in range(1, max_retries + 1):
            try:
                self.write({
                    'processing_log': (self.processing_log or '') +
                                      f'เรียก Gemini (ครั้งที่ {attempt})...\n',
                })
                gemini_result = self._call_gemini(api_key, model_name, prompt)

                self.write({
                    'processing_log': (self.processing_log or '') +
                                      f'Gemini ตอบกลับ {len(gemini_result)} chars\n',
                })

                reconstructed, required_fields_list = self._parse_gemini_response(gemini_result)
                if reconstructed:
                    break  # success
                last_error = 'AI ไม่พบช่องกรอกในเอกสาร'
            except Exception as e:
                last_error = str(e)
                _logger.warning('Gemini attempt %d failed: %s', attempt, e)
                self.write({
                    'processing_log': (self.processing_log or '') +
                                      f'ครั้งที่ {attempt} ไม่สำเร็จ: {e}\n',
                })

        if not reconstructed:
            raise UserError(
                f'AI วิเคราะห์ไม่สำเร็จ หลังลอง {max_retries} ครั้ง\n'
                f'ข้อผิดพลาดล่าสุด: {last_error}'
            )

        # 4) Create processed docx with {{placeholders}}
        processed_bytes = self._apply_reconstructed_to_docx(docx_bytes, reconstructed)

        # 5) Save processed docx to file system for template-fill
        safe_code = re.sub(r'[^\w\-/]', '', self.code or 'template')
        file_path = f'docx-masters/{safe_code}.docx'
        self._save_to_template_dir(file_path, processed_bytes)

        # 6) Update record
        safe_name = re.sub(r'[^\w\s\-]', '', self.name or 'template').strip()
        self.write({
            'processed_docx': base64.b64encode(processed_bytes),
            'processed_docx_filename': f'{safe_name}_template.docx',
            'required_fields': json.dumps(required_fields_list, ensure_ascii=False, indent=2),
            'template_file_path': file_path,
            'processing_state': 'ready',
            'processing_log': (self.processing_log or '') +
                              f'สำเร็จ! แก้ไข {len(reconstructed)} ย่อหน้า\n'
                              f'สร้าง {len(required_fields_list)} fields\n'
                              f'บันทึกไฟล์ที่ {file_path}\n',
        })

    def _build_analysis_prompt(self, doc_text):
        """Build the prompt for Gemini to analyze a Thai legal document."""
        return f"""คุณเป็นผู้เชี่ยวชาญด้านเอกสารกฎหมายไทย วิเคราะห์เอกสารต่อไปนี้และระบุช่องที่ต้องกรอกข้อมูล

เอกสาร (แต่ละบรรทัดมี [index] กำกับ):
---
{doc_text}
---

งานของคุณ:
1. หาส่วนที่เป็นข้อมูลที่ต้องกรอก เช่น ชื่อ, วันที่, ที่อยู่, จำนวนเงิน, เลขบัตรประชาชน ฯลฯ
2. หาส่วนที่เป็นจุด (....) หรือเส้น (____) สำหรับกรอก
3. สร้าง placeholder ภาษาไทยที่เข้าใจง่าย

สำคัญมาก: สร้างผลลัพธ์เป็น "reconstructed_paragraphs" — เขียนทั้งย่อหน้าใหม่โดยแทนช่องกรอกด้วย {{{{placeholder}}}}
ห้าม replace ทีละส่วน ให้เขียนทั้งย่อหน้าใหม่ เพื่อไม่ให้ placeholder สับสน

ตอบเป็น JSON เท่านั้น:
```json
{{
  "reconstructed_paragraphs": [
    {{
      "paragraph_index": 1,
      "new_text": "เขียนที่ {{{{สถานที่เขียน}}}} วันที่ {{{{วันที่}}}} เดือน {{{{เดือน}}}} พ.ศ. {{{{ปีพศ}}}}"
    }}
  ],
  "required_fields": [
    {{
      "name": "สถานที่เขียน",
      "label": "สถานที่เขียน",
      "type": "text",
      "required": true
    }}
  ]
}}
```

กฎ:
- placeholder ใช้ภาษาไทย เข้าใจง่าย เช่น {{{{ชื่อผู้มอบอำนาจ}}}}, {{{{วันที่}}}}, {{{{เลขบัตรประชาชน}}}}
- ทุก placeholder ต้อง unique — ห้ามซ้ำกัน ถ้ามีหลายคน ให้ใส่เลข เช่น {{{{ชื่อผู้เช่า1}}}}, {{{{ชื่อผู้เช่า2}}}}
- ถ้าเป็นจุด (....) หรือเส้น (____) หรือช่องว่างสำหรับกรอก ให้แทนด้วย placeholder ที่เหมาะสมกับบริบท
- ข้อความที่ไม่ใช่ช่องกรอก ให้คงเดิมทุกตัวอักษร
- type: "text" สำหรับชื่อ/ที่อยู่, "number" สำหรับตัวเลข, "date" สำหรับวันที่, "textarea" สำหรับข้อความยาว
- required: true สำหรับฟิลด์จำเป็น
- เรียงลำดับ required_fields ตามลำดับที่ปรากฏในเอกสาร
- paragraph_index ต้องตรงกับ [index] ของเอกสาร
- ตอบเฉพาะ JSON object เท่านั้น ห้ามมีข้อความอื่นนอกเหนือจาก JSON
- ห้ามใช้ trailing comma ใน JSON"""

    def _call_gemini(self, api_key, model_name, prompt):
        """Call Gemini REST API directly with JSON mode."""
        import requests

        url = f'https://generativelanguage.googleapis.com/v1beta/models/{model_name}:generateContent'
        resp = requests.post(
            url,
            params={'key': api_key},
            json={
                'contents': [{'parts': [{'text': prompt}]}],
                'generationConfig': {
                    'temperature': 0.1,
                    'maxOutputTokens': 65536,
                    'responseMimeType': 'application/json',
                },
            },
            timeout=180,
        )
        resp.raise_for_status()
        data = resp.json()

        # Extract text from response
        candidates = data.get('candidates', [])
        if not candidates:
            _logger.error('Gemini response: %s', json.dumps(data, ensure_ascii=False)[:1000])
            raise UserError('Gemini ไม่ส่งผลลัพธ์กลับมา')

        # Check finish reason
        finish_reason = candidates[0].get('finishReason', '')
        if finish_reason == 'MAX_TOKENS':
            _logger.warning('Gemini response truncated (MAX_TOKENS)')

        parts = candidates[0].get('content', {}).get('parts', [])
        return ''.join(p.get('text', '') for p in parts)

    def _repair_json(self, json_str):
        """Attempt to repair common JSON issues from LLM output."""
        # Remove markdown code fences if present
        json_str = re.sub(r'^```json\s*', '', json_str.strip())
        json_str = re.sub(r'\s*```$', '', json_str.strip())

        # Try direct parse first
        try:
            return json.loads(json_str)
        except json.JSONDecodeError:
            pass

        # Fix trailing commas before } or ]
        fixed = re.sub(r',\s*([}\]])', r'\1', json_str)
        try:
            return json.loads(fixed)
        except json.JSONDecodeError:
            pass

        # If JSON is truncated, try to close it
        truncated = fixed.rstrip()
        # Count open/close braces and brackets
        open_braces = truncated.count('{') - truncated.count('}')
        open_brackets = truncated.count('[') - truncated.count(']')
        truncated += ']' * max(0, open_brackets) + '}' * max(0, open_braces)
        try:
            return json.loads(truncated)
        except json.JSONDecodeError:
            pass

        return None

    def _parse_gemini_response(self, response_text):
        """Parse Gemini's JSON response into reconstructed_paragraphs and required_fields."""
        # Extract JSON from response (may be wrapped in ```json ... ```)
        json_match = re.search(r'```json\s*(.*?)\s*```', response_text, re.DOTALL)
        if json_match:
            json_str = json_match.group(1)
        else:
            json_match = re.search(r'\{[\s\S]*\}', response_text)
            if json_match:
                json_str = json_match.group(0)
            else:
                json_str = response_text

        # Try repair
        data = self._repair_json(json_str)
        if data is None:
            _logger.error('Failed to parse Gemini JSON:\n%s', response_text[:2000])
            raise UserError(
                f'ไม่สามารถ parse JSON จาก AI\n'
                f'ข้อความดิบ (500 ตัวแรก):\n{response_text[:500]}'
            )

        reconstructed = data.get('reconstructed_paragraphs', [])
        required_fields = data.get('required_fields', [])
        return reconstructed, required_fields

    def _apply_reconstructed_to_docx(self, docx_bytes, reconstructed):
        """Replace entire paragraphs in docx using AI's reconstructed text."""
        from docx import Document as DocxDoc

        doc = DocxDoc(io.BytesIO(docx_bytes))

        # Build index map: paragraph_index -> new_text
        recon_map = {}
        for item in reconstructed:
            idx = item.get('paragraph_index')
            new_text = item.get('new_text', '')
            if idx is not None and new_text:
                recon_map[idx] = new_text

        # Apply: for each matching paragraph, put new_text in first run, clear rest
        for i, para in enumerate(doc.paragraphs):
            if i in recon_map:
                new_text = recon_map[i]
                if para.runs:
                    # Preserve formatting of first run, replace text
                    para.runs[0].text = new_text
                    for run in para.runs[1:]:
                        run.text = ''
                else:
                    # No runs — add one
                    para.add_run(new_text)

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    def _save_to_template_dir(self, file_path, docx_bytes):
        """Save processed docx to /mnt/templates/ directory."""
        base_path = '/mnt/templates'
        full_path = os.path.join(base_path, file_path)
        os.makedirs(os.path.dirname(full_path), exist_ok=True)
        with open(full_path, 'wb') as f:
            f.write(docx_bytes)
        _logger.info('Saved processed template to %s (%d bytes)',
                     full_path, len(docx_bytes))

    def action_preview_fields(self):
        """Button: preview required_fields in a popup."""
        self.ensure_one()
        return {
            'type': 'ir.actions.act_window',
            'name': f'Fields: {self.name}',
            'res_model': 'legal.document.template',
            'res_id': self.id,
            'view_mode': 'form',
            'target': 'new',
        }

    def action_reset_processing(self):
        """Reset processing state — clear AI results and log."""
        self.ensure_one()
        self.write({
            'processing_state': 'uploaded' if self.master_docx else 'none',
            'processing_log': '',
            'processed_docx': False,
            'processed_docx_filename': False,
            'required_fields': False,
            'template_file_path': False,
        })
